"""Word export for publication-style research papers."""
from __future__ import annotations

import io
from datetime import datetime

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from utils.export_utils import markdown_soup, references_table_data, sanitize_text, table_matrix


def _set_document_defaults(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = document.styles
    styles["Normal"].font.name = "Times New Roman"
    styles["Normal"].font.size = Pt(11)


def _add_header_watermark(document: Document, watermark_text: str) -> None:
    if not watermark_text:
        return

    for section in document.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(watermark_text)
        run.font.name = "Calibri"
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(190, 190, 190)
        run.italic = True


def _add_paragraph(document: Document, text: str, style: str | None = None, italic: bool = False) -> None:
    paragraph = document.add_paragraph(style=style)
    run = paragraph.add_run(sanitize_text(text))
    run.italic = italic


def _add_table(document: Document, data: list[list[str]]) -> None:
    if not data:
        return
    table = document.add_table(rows=len(data), cols=len(data[0]))
    table.style = "Table Grid"
    for row_index, row in enumerate(data):
        for col_index, value in enumerate(row):
            table.cell(row_index, col_index).text = sanitize_text(value)


def generate_docx(
    query: str,
    answer: str,
    sources: list[dict],
    confidence: float = 0.0,
    timestamp: str | None = None,
    watermark_text: str | None = None,
) -> bytes:
    document = Document()
    _set_document_defaults(document)
    _add_header_watermark(document, watermark_text or "")

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run(sanitize_text(query))
    title_run.bold = True
    title_run.font.name = "Times New Roman"
    title_run.font.size = Pt(18)

    meta = document.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        f"DoraEngine Research Paper | Confidence {int(confidence * 100)}% | {timestamp or datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
    )
    meta_run.font.name = "Calibri"
    meta_run.font.size = Pt(10)
    meta_run.font.color.rgb = RGBColor(90, 90, 90)

    soup = markdown_soup(answer)
    root = soup.find("div")
    for tag in root.children:
        if not getattr(tag, "name", None):
            continue
        name = tag.name.lower()
        if name == "h1":
            _add_paragraph(document, tag.get_text(" ", strip=True), "Heading 1")
        elif name == "h2":
            _add_paragraph(document, tag.get_text(" ", strip=True), "Heading 2")
        elif name == "h3":
            _add_paragraph(document, tag.get_text(" ", strip=True), "Heading 3")
        elif name == "p":
            _add_paragraph(document, tag.get_text(" ", strip=True))
        elif name in {"ul", "ol"}:
            for item in tag.find_all("li", recursive=False):
                style = "List Bullet" if name == "ul" else "List Number"
                _add_paragraph(document, item.get_text(" ", strip=True), style)
        elif name == "table":
            _add_table(document, table_matrix(tag))
        elif name == "blockquote":
            _add_paragraph(document, tag.get_text(" ", strip=True), italic=True)

    if sources:
        _add_paragraph(document, "References", "Heading 2")
        _add_table(document, references_table_data(sources))

    output = io.BytesIO()
    document.save(output)
    return output.getvalue()
